#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word -> Platform Question Importer

Buyruqlar:
  python word_import.py shablon
      savollar_shablon.docx yaratadi

  python word_import.py korish --fayl savollar.docx
      yuklashdan oldin ko'rib chiqish

  python word_import.py yuklash --fayl savollar.docx --assignment UUID
      API ga yuklash
"""

import sys
import io
import json
import argparse
import requests
from pathlib import Path

# Windows konsolida UTF-8
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx o'rnatilmagan. Quyidagini bajaring:")
    print("  pip install python-docx requests")
    sys.exit(1)

# ─────────────────────────────────────────────────────────────
#  SOZLAMALAR
# ─────────────────────────────────────────────────────────────
API_URL       = "http://localhost:8000/api"
ASSIGNMENT_ID = ""   # default assignment UUID (ixtiyoriy)

# ─────────────────────────────────────────────────────────────
#  YORDAMCHI
# ─────────────────────────────────────────────────────────────
def _cell(table, row, col):
    try:
        return table.cell(row, col).text.strip()
    except IndexError:
        return ""

def _key(table, row):
    return _cell(table, row, 0).upper().strip()

def _val(table, row):
    return _cell(table, row, 1)


# ─────────────────────────────────────────────────────────────
#  PARSERS
# ─────────────────────────────────────────────────────────────

def parse_variantli(table):
    """
    | VARIANTLI |                              |
    | SAVOL     | Savol matni                  |
    | A         | Variant 1                    |
    | B*        | To'g'ri variant  <- *        |
    | C         | Variant 3                    |
    | BALL      | 5                            |
    | IZOH      | Tushuntirish (ixtiyoriy)     |
    """
    savol   = ""
    options = []
    correct = ""
    points  = 1
    izoh    = ""

    for i in range(1, len(table.rows)):
        k = _key(table, i)
        v = _val(table, i)

        if k == "SAVOL":
            savol = v
        elif k in ("A","B","C","D","E","F","A*","B*","C*","D*","E*","F*"):
            is_correct = k.endswith("*")
            if v:
                options.append(v)
                if is_correct:
                    correct = v
        elif k == "IZOH":
            izoh = v
        elif k == "BALL":
            try:
                points = int(v)
            except ValueError:
                points = 1

    if not savol:
        return None, "SAVOL bo'sh"
    if len(options) < 2:
        return None, "Kamida 2 ta variant kerak"
    if not correct:
        return None, "To'g'ri javob ko'rsatilmagan (variant yoniga * qo'ying, masalan: B*)"

    q = {
        "question_text": savol,
        "question_data": {"options": options},
        "correct_answer": correct,
        "points": points,
    }
    if izoh:
        q["explanation"] = izoh
    return q, None


def parse_ha_yoq(table):
    """
    | HA_YOQ |                           |
    | SAVOL  | Python interpreted tildir |
    | JAVOB  | HA                        |
    | BALL   | 3                         |
    | IZOH   | (ixtiyoriy)               |
    """
    savol  = ""
    javob  = ""
    points = 1
    izoh   = ""

    for i in range(1, len(table.rows)):
        k = _key(table, i)
        v = _val(table, i)

        if k == "SAVOL":
            savol = v
        elif k == "JAVOB":
            raw = v.upper().strip()
            if raw in ("HA", "TO'G'RI", "TRUE", "1", "+"):
                javob = "true"
            elif raw in ("YO'Q", "YOQ", "NOTO'G'RI", "FALSE", "0", "-"):
                javob = "false"
        elif k == "IZOH":
            izoh = v
        elif k == "BALL":
            try:
                points = int(v)
            except ValueError:
                points = 1

    if not savol:
        return None, "SAVOL bo'sh"
    if not javob:
        return None, "JAVOB ko'rsatilmagan (HA yoki YO'Q yozing)"

    q = {
        "question_text": savol,
        "question_data": {"statement": savol},
        "correct_answer": javob,
        "points": points,
    }
    if izoh:
        q["explanation"] = izoh
    return q, None


def parse_bosh(table):
    """
    | BO'SH  |                                        |
    | MATN   | Python ___ tilida yaratilgan           |
    | JAVOB  | dasturlash                             |
    | BALL   | 4                                      |
    | IZOH   | (ixtiyoriy)                            |

    Bo'sh joylar uchun ___  (3 ta pastki chiziq).
    """
    matn   = ""
    javob  = ""
    points = 1
    izoh   = ""

    for i in range(1, len(table.rows)):
        k = _key(table, i)
        v = _val(table, i)

        if k == "MATN":
            matn = v
        elif k in ("JAVOB", "JAVOB1"):
            javob = v
        elif k == "IZOH":
            izoh = v
        elif k == "BALL":
            try:
                points = int(v)
            except ValueError:
                points = 1

    if not matn:
        return None, "MATN bo'sh"
    if not javob:
        return None, "JAVOB ko'rsatilmagan"
    if "___" not in matn:
        return None, "MATNda ___ bo'sh joy belgisi yo'q"

    q = {
        "question_text": matn,
        "question_data": {"text": matn, "blanks": matn.count("___")},
        "correct_answer": javob,
        "points": points,
    }
    if izoh:
        q["explanation"] = izoh
    return q, None


def parse_moslashtirish(table):
    """
    | MOSLASHTIRISH |                  |
    | SAVOL         | Mos keltiring    |   <- ixtiyoriy
    | CPU           | Protsessor       |   <- juftliklar
    | RAM           | Operativ xotira  |
    | BALL          | 6                |
    | IZOH          | (ixtiyoriy)      |
    """
    RESERVED = {"MOSLASHTIRISH", "MATCHING", "SAVOL", "BALL", "IZOH", ""}

    savol       = "Quyidagilarni mos keltiring:"
    pairs       = []
    correct_map = {}
    points      = 1
    izoh        = ""

    for i in range(1, len(table.rows)):
        k             = _key(table, i)
        v             = _val(table, i)
        original_left = _cell(table, i, 0).strip()

        if k == "SAVOL":
            savol = v or savol
        elif k == "IZOH":
            izoh = v
        elif k == "BALL":
            try:
                points = int(v)
            except ValueError:
                points = 1
        elif k not in RESERVED and original_left and v:
            pairs.append({"left": original_left, "right": v})
            correct_map[original_left] = v

    if len(pairs) < 2:
        return None, "Kamida 2 ta juftlik kerak"

    q = {
        "question_text": savol,
        "question_data": {"pairs": pairs},
        "correct_answer": correct_map,
        "points": points,
    }
    if izoh:
        q["explanation"] = izoh
    return q, None


# ─────────────────────────────────────────────────────────────
#  TUR -> PARSER XARITASI
# ─────────────────────────────────────────────────────────────
TYPE_PARSERS = {
    "VARIANTLI":           parse_variantli,
    "KO'P TANLOV":         parse_variantli,
    "MULTIPLE CHOICE":     parse_variantli,
    "HA_YOQ":              parse_ha_yoq,
    "HA/YO'Q":             parse_ha_yoq,
    "HA YO'Q":             parse_ha_yoq,
    "TRUE FALSE":          parse_ha_yoq,
    "TO'G'RI/NOTO'G'RI":  parse_ha_yoq,
    "BO'SH":               parse_bosh,
    "BO'SH TO'LDIRISH":    parse_bosh,
    "FILLING":             parse_bosh,
    "FILL BLANK":          parse_bosh,
    "MOSLASHTIRISH":       parse_moslashtirish,
    "MATCHING":            parse_moslashtirish,
}


def parse_document(docx_path):
    """Word hujjatidagi barcha jadvallarni parse qilish."""
    doc       = Document(str(docx_path))
    questions = []
    errors    = []

    for idx, table in enumerate(doc.tables, 1):
        if len(table.rows) < 2:
            continue

        type_raw = _cell(table, 0, 0).strip().upper()
        parser   = TYPE_PARSERS.get(type_raw)

        if parser is None:
            errors.append({
                "jadval": idx,
                "xato": f"Noma'lum tur '{type_raw}'. "
                        f"Mavjud turlar: {', '.join(sorted(set(TYPE_PARSERS)))}"
            })
            continue

        q, err = parser(table)
        if err:
            errors.append({"jadval": idx, "tur": type_raw, "xato": err})
            continue

        q["order_index"] = len(questions) + 1
        questions.append(q)

    return questions, errors


# ─────────────────────────────────────────────────────────────
#  API YUKLOVCHI
# ─────────────────────────────────────────────────────────────
def upload_questions(questions, assignment_id, api_url):
    print(f"\nYuklanmoqda -> {api_url}/questions/")
    print(f"Assignment  -> {assignment_id}\n")

    success = 0
    for q in questions:
        payload = {**q, "assignment": assignment_id}
        try:
            r = requests.post(f"{api_url}/questions/", json=payload, timeout=10)
            if r.status_code == 201:
                success += 1
                print(f"  [OK] [{q['order_index']:02d}] {q['question_text'][:60]}")
            else:
                print(f"  [!!] [{q['order_index']:02d}] HTTP {r.status_code}: {r.text[:100]}")
        except requests.ConnectionError:
            print(f"  [!!] Server bilan aloqa yo'q -> {api_url}")
            break
        except requests.RequestException as e:
            print(f"  [!!] Tarmoq xatosi: {e}")

    print(f"\nNatija: {success} / {len(questions)} savol yuklandi.")


# ─────────────────────────────────────────────────────────────
#  SHABLON YARATISH
# ─────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_question_table(doc, rows, header_hex="4F46E5", key_hex="E0E7FF"):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"

    for i, (k, v) in enumerate(rows):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)

        c0.width = Cm(4.5)
        c1.width = Cm(12.5)

        c0.text = k
        c1.text = v

        p0 = c0.paragraphs[0]
        if p0.runs:
            p0.runs[0].bold = True
            p0.runs[0].font.size = Pt(10)

        if i == 0:
            _set_cell_bg(c0, header_hex)
            _set_cell_bg(c1, header_hex)
            for cell in (c0, c1):
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
        else:
            _set_cell_bg(c0, key_hex)

    doc.add_paragraph()


def create_template(output_path="savollar_shablon.docx"):
    doc = Document()

    section = doc.sections[0]
    section.page_width   = Cm(29.7)
    section.page_height  = Cm(21)
    section.left_margin  = Cm(2)
    section.right_margin = Cm(2)

    doc.add_heading("Savollar Shabloni", 0)
    doc.add_paragraph(
        "Qoidalar:\n"
        "  1. Har bir savol = alohida jadval\n"
        "  2. Birinchi qator - savol turi (o'zgartirmang)\n"
        "  3. BALL qatoriga faqat son yozing\n"
        "  4. IZOH qatori ixtiyoriy\n"
        "  5. Formulalar: $x^2 + 1 = 0$ yoki $$\\frac{a}{b}$$"
    )
    doc.add_paragraph()

    # 1. VARIANTLI
    doc.add_heading("1. Ko'p tanlovli (VARIANTLI)", level=2)
    doc.add_paragraph(
        "To'g'ri variant harfi yoniga * qo'ying. Masalan: B*"
    ).italic = True
    _add_question_table(doc, [
        ("VARIANTLI",  ""),
        ("SAVOL",      "Qaysi dasturlash tili eng ko'p qo'llaniladi?"),
        ("A",          "Java"),
        ("B*",         "Python    <- to'g'ri javob"),
        ("C",          "C++"),
        ("D",          "Ruby"),
        ("BALL",       "5"),
        ("IZOH",       "Python data science sohasida keng qo'llaniladi.  (ixtiyoriy)"),
    ], header_hex="4F46E5", key_hex="E0E7FF")

    # 2. HA / YO'Q
    doc.add_heading("2. Ha / Yo'q (HA_YOQ)", level=2)
    doc.add_paragraph(
        "JAVOB: HA yoki YO'Q"
    ).italic = True
    _add_question_table(doc, [
        ("HA_YOQ",  ""),
        ("SAVOL",   "Python - interpreted (tarjima qilib bajariladigan) tildir."),
        ("JAVOB",   "HA"),
        ("BALL",    "3"),
        ("IZOH",    "Python CPython interpretatori orqali ishlaydi.  (ixtiyoriy)"),
    ], header_hex="059669", key_hex="D1FAE5")

    # 3. BO'SH TO'LDIRISH
    doc.add_heading("3. Bo'sh to'ldirish (BO'SH)", level=2)
    doc.add_paragraph(
        "Bo'sh joylar uchun ___ (uch pastki chiziq) ishlating."
    ).italic = True
    _add_question_table(doc, [
        ("BO'SH",  ""),
        ("MATN",   "Python ___ tilida yaratilgan bo'lib, ___ yili chiqarilgan."),
        ("JAVOB",  "dasturlash, 1991"),
        ("BALL",   "4"),
        ("IZOH",   "Guido van Rossum tomonidan yaratilgan.  (ixtiyoriy)"),
    ], header_hex="D97706", key_hex="FEF3C7")
    doc.add_paragraph(
        "Bir nechta bo'sh joy bo'lsa JAVOB da vergul bilan ajrating."
    ).runs[0].font.size = Pt(9)

    # 4. MOSLASHTIRISH
    doc.add_heading("4. Moslashtirish (MOSLASHTIRISH)", level=2)
    doc.add_paragraph(
        "Chap ustun: element  |  O'ng ustun: mos juft\n"
        "SAVOL, BALL, IZOH dan boshqa barcha qatorlar juftlik sifatida o'qiladi."
    ).italic = True
    _add_question_table(doc, [
        ("MOSLASHTIRISH",  ""),
        ("SAVOL",          "Quyidagi terminlarni mos ta'rif bilan jufting:"),
        ("CPU",            "Markaziy protsessor"),
        ("RAM",            "Operativ xotira"),
        ("SSD",            "Qattiq holat disk"),
        ("GPU",            "Grafik protsessor"),
        ("BALL",           "8"),
        ("IZOH",           "Kompyuter arxitekturasi asoslari.  (ixtiyoriy)"),
    ], header_hex="DC2626", key_hex="FEE2E2")

    # 2-bet: bo'sh savollar uchun
    doc.add_page_break()
    doc.add_heading("Savollar (shu yerdan boshlang)", level=1)
    doc.add_paragraph(
        "Yuqoridagi namunalardan birini nusxa olib, "
        "mazmunini o'zgartirib to'ldiring."
    )

    doc.save(str(output_path))

    print(f"")
    print(f"[OK] Shablon yaratildi -> {output_path}")
    print(f"")
    print(f"Keyingi qadamlar:")
    print(f"  1. {output_path} ni Word da oching")
    print(f"  2. 2-betga o'ting")
    print(f"  3. Namunani nusxa olib, o'zingiznikini to'ldiring")
    print(f"  4. .docx formatida saqlang")
    print(f"  5. python word_import.py korish --fayl {output_path}")


# ─────────────────────────────────────────────────────────────
#  CLI
# ─────────────────────────────────────────────────────────────
def main():
    p = argparse.ArgumentParser(description="Word -> Platform Question Importer")
    sub = p.add_subparsers(dest="cmd", required=True)

    # shablon
    s = sub.add_parser("shablon", help="Word shablon yaratish")
    s.add_argument("--chiqish", "-o", default="savollar_shablon.docx")

    # korish
    k = sub.add_parser("korish", help="Ko'rib chiqish (yuklamasdan)")
    k.add_argument("--fayl", "-f", required=True)

    # yuklash
    y = sub.add_parser("yuklash", help="API ga yuklash")
    y.add_argument("--fayl",       "-f", required=True)
    y.add_argument("--assignment", "-a", default=ASSIGNMENT_ID)
    y.add_argument("--api",              default=API_URL)

    args = p.parse_args()

    if args.cmd == "shablon":
        create_template(args.chiqish)
        return

    # korish / yuklash
    path = Path(args.fayl)
    if not path.exists():
        print(f"Xato: fayl topilmadi -> {path}")
        sys.exit(1)
    if path.suffix.lower() != ".docx":
        print("Xato: faqat .docx format qabul qilinadi")
        sys.exit(1)

    print(f"\nO'qilmoqda: {path}")
    questions, errors = parse_document(path)

    if errors:
        print(f"\n[OGOHLANTIRISH] {len(errors)} ta muammo:")
        for e in errors:
            jadval = e.get("jadval", "?")
            tur    = e.get("tur", "")
            label  = f"[{tur}] " if tur else ""
            print(f"  Jadval #{jadval} {label}-> {e['xato']}")

    SEP = "-" * 62
    print(f"\n{SEP}")
    print(f"Topildi: {len(questions)} ta savol")
    print(SEP)

    if not questions:
        print("Yuklanadigan savol yo'q.")
        return

    TUR_LABEL = {
        "options":   "Variantli      ",
        "statement": "Ha/Yo'q        ",
        "text":      "Bo'sh toldirish",
        "pairs":     "Moslashtirish  ",
    }
    for q in questions:
        first_key = next(iter(q["question_data"]), "")
        tur_label  = TUR_LABEL.get(first_key, "Noma'lum       ")
        savol_qisq = q["question_text"][:52]
        print(f"  [{q['order_index']:02d}] {tur_label} | {q['points']}b | {savol_qisq}")

    if args.cmd == "korish":
        print(f"\n{SEP}")
        print("JSON natijasi:")
        print(json.dumps(questions, ensure_ascii=False, indent=2))
        print(f"\nYuklash uchun:")
        print(f"  python word_import.py yuklash --fayl {path} --assignment <UUID>")
        return

    # yuklash
    if not args.assignment:
        print("\nXato: --assignment <UUID> parametri kerak")
        sys.exit(1)

    print(SEP)
    yn = input("Yuklashni boshlashni tasdiqlaysizmi? (ha/yo'q): ").strip().lower()
    if yn not in ("ha", "h", "yes", "y"):
        print("Bekor qilindi.")
        return

    upload_questions(questions, args.assignment, args.api)


if __name__ == "__main__":
    main()
